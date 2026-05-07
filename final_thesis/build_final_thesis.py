# -*- coding: utf-8 -*-
"""
论文输出：Markdown + Word + PDF（文件名不含「终版」后缀）
- 正文来自 thesis_body_content.md
- 精选随文分析图：第3章场地分析、第5章方案与功能、第6章技术图纸（不含效果图文件夹）
- 不使用 OCR；每张图配「读图分析」段落（结合开题与图纸内容撰写）
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = Path(__file__).resolve().parent
DRAWINGS = ROOT / "drawings"
MD_NAME = "彭州市水街茶咖啡厅室内设计_李青桐.md"
DOCX_NAME = "彭州市水街茶咖啡厅室内设计_李青桐.docx"
PDF_NAME = "彭州市水街茶咖啡厅室内设计_李青桐.pdf"


def set_run_font(
    run,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size_pt: float = 12,
    bold: bool | None = None,
    *,
    color: RGBColor | None = RGBColor(0, 0, 0),
):
    """正文按规范：中文 East Asia 字体 + 西文/数字 Times New Roman；默认黑色。"""
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), east_asia)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold


_MD_BOLD = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)


def iter_md_bold_chunks(text: str):
    """将段落中的 `**粗体**` 拆成 (是否粗体, 片段) 序列；无标记时整段为普通文本。"""
    pos = 0
    for m in _MD_BOLD.finditer(text):
        if m.start() > pos:
            yield False, text[pos : m.start()]
        yield True, m.group(1)
        pos = m.end()
    if pos < len(text):
        yield False, text[pos:]


def paragraph_add_body_runs(paragraph, text: str, *, body_pt: float):
    """写入正文段落：支持 Markdown 粗体标记，不保留星号。"""
    for is_bold, chunk in iter_md_bold_chunks(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        set_run_font(
            run,
            east_asia="宋体",
            ascii_font="Times New Roman",
            size_pt=body_pt,
            bold=True if is_bold else False,
        )


def _style_rfonts(style):
    """为样式设置 eastAsia 字体（兼容部分环境下 rPr 缺失）。"""
    el = style._element
    rPr = el.get_or_add_rPr()
    return rPr.get_or_add_rFonts()


def configure_document_styles(doc: Document) -> None:
    """
    按《四川农业大学本科生毕业论文（设计）撰写规范（教发〔2016〕10号）》OCR 摘录：
    - 正文：宋体小四（12pt），英文数字 Times New Roman，行距 1.5 倍，首行缩进 2 字符，标准字距，黑色；
    - 一级标题（章）：黑体小三号（15pt）；二级（节）：黑体四号（14pt）；
    - 图题：黑体五号（10.5pt），置于图下方；参考文献条目：五号（10.5pt）。
    """
    styles = doc.styles

    n = styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(12)
    n.font.color.rgb = RGBColor(0, 0, 0)
    _style_rfonts(n).set(qn("w:eastAsia"), "宋体")
    np = n.paragraph_format
    np.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    np.line_spacing = 1.5
    np.first_line_indent = Cm(0.74)
    np.space_before = Pt(0)
    np.space_after = Pt(0)

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    _style_rfonts(h1).set(qn("w:eastAsia"), "黑体")
    h1p = h1.paragraph_format
    h1p.first_line_indent = Pt(0)
    h1p.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1p.line_spacing = 1.5
    h1p.space_before = Pt(12)
    h1p.space_after = Pt(6)
    h1p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1p.page_break_before = False

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    _style_rfonts(h2).set(qn("w:eastAsia"), "黑体")
    h2p = h2.paragraph_format
    h2p.first_line_indent = Pt(0)
    h2p.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2p.line_spacing = 1.5
    h2p.space_before = Pt(6)
    h2p.space_after = Pt(3)
    h2p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def sorted_images_in(folder_name: str) -> list[Path]:
    d = DRAWINGS / folder_name
    if not d.is_dir():
        return []
    return sorted([p for p in d.iterdir() if p.suffix.lower() in (".png", ".jpg", ".jpeg")])


def figures_in_order(folder_name: str, filenames: list[str]) -> list[Path]:
    """按给定文件名顺序取图（仅保留磁盘上存在的项），用于精选重点图纸。"""
    d = DRAWINGS / folder_name
    out: list[Path] = []
    if not d.is_dir():
        return out
    for name in filenames:
        p = d / name
        if p.is_file():
            out.append(p)
    return out


def figures_from_tuples(items: list[tuple[str, str]]) -> list[Path]:
    """按 (子文件夹名, 文件名) 列表依次收集路径。"""
    out: list[Path] = []
    for folder_name, fname in items:
        p = DRAWINGS / folder_name / fname
        if p.is_file():
            out.append(p)
    return out


def rel_md_path(fig: Path) -> str:
    return "../" + str(fig.relative_to(ROOT)).replace("\\", "/")


def fig_key(fig: Path) -> str:
    return fig.relative_to(ROOT).as_posix()


def narrative_for_figure(fig: Path) -> str:
    """基于对图纸的阅读（含实地读图与图面信息）撰写的分析段落，不用 OCR。"""
    n = fig.name
    parent = fig.parent.name

    # —— 前期与综合分析 ——
    if n == "区位分析图.png":
        return (
            "该图为灰度大比例区位底图，在四川省域与成都平原西北部语境中标示彭州与周边城镇组群的关系。"
            "深色范围强调项目所处片区在区域交通走廊中的位置，道路网络呈向心—联络式结构，说明水街作为老城更新节点可承接外向游客与内向通勤。"
            "右上角省域缩略与右下实景鸟瞰镶嵌，使宏观尺度与街区肌理同时可读，为后续论证“文旅商业+本地生活”双客群提供空间依据。"
        )
    if n == "区位分析图1.png":
        return (
            "本图为区位分析的补充页，从更近尺度展开周边道路、用地与重要设施关系，便于识别主要来向与潜在人流走廊。"
            "与上一张配合阅读，可判断茶咖入口朝向与外部广场、步行街的衔接策略，避免仅停留在概念层面的“地段好”。"
        )
    if n == "场地交通流线分析.png":
        return (
            "图纸以轴测方式呈现水街片区整体形态：中央水体与环形、月牙形步行桥构成强视觉核心，彩色流线区分公共、文化与酒店等不同人群路径。"
            "地面一层实线与二层虚线表达立体漫游系统，说明设计关注“慢行体验”与多点进入的组织方式。"
            "传统院落组群与现代体块并置，流线在滨水与院落之间切换，可为室内设计提供“外部叙事—内部节奏”的连续性参考。"
        )
    if n == "场地功能分区分析.png":
        return (
            "功能分区图以颜色块与图例界定民俗商业、文化设施、市民活动与建筑改造等区域，显示茶咖所处的外部功能环境并非单一商业街。"
            "阅读重点在于：相邻功能是否带来白天与夜晚人流差异、是否存在文化参观后的休憩需求，以及外部活动是否会对室内声学与照明提出联动要求。"
        )
    if n == "建筑结构分析图.png":
        return (
            "结构分析图揭示主体体量、屋面坡度与楼梯等竖向构件关系，是后续墙体改造与吊顶设计的上位条件。"
            "阅读时应关注柱网间距、梁板传力方向与可拆除填充墙范围，避免平面自由化导致结构风险。"
        )
    if n == "冬夏两季日照分析.png":
        return (
            "日照分析以简化的体块模型叠加冬夏太阳轨迹弧线：夏季弧高、冬季弧低，标示日出日落与正午方位。"
            "读图可知设计意图在于利用南向大面采光并在夏季通过挑檐、屋檐或遮阳手段削弱西晒与过热风险，冬季则争取更低太阳高度角下的室内进深采光。"
            "该结论直接支持室内窗帘选型、吊顶反檐深度与座位靠窗区的热舒适处理。"
        )
    if n == "通风分析图.png":
        return (
            "通风分析通常结合主导风向与开口位置，提示可开启扇、空调风口与气味控制的基本策略。"
            "对茶咖而言，吧台与冲泡区的气味溢散、围炉区的热烟上升路径，都应在通风逻辑中被预判并与吊顶设备带协同。"
        )

    if n == "03_一层建筑墙体改造图.png":
        return (
            "墙体改造图标示可拆填充墙、保留承重墙与新增轻质隔断的位置，是平面与结构条件之间的接口图纸。"
            "阅读时应核对楼梯、卫生间与后场是否仍满足疏散净宽，并预判改造后吊顶与机电路由的走向。"
        )
    if n == "04_二层建筑墙体改造图.png":
        return (
            "二层墙体改造与一层策略应连贯：体验区开敞与雅间围合的边界需在此图上闭合，避免仅在一层平面完成分区而竖向错位。"
        )
    if n == "05_一层地面铺装图.png":
        return (
            "铺装图以图案与分缝表达主通道、停留区与展示区的地面差异，材料交接与找坡方向应结合排水点位阅读。"
        )
    if n == "06_二层地面铺装图.png":
        return (
            "二层铺装需与家具固定点、地插及围炉区防护协调；读图重点为铺装分缝是否穿越主要动线造成绊脚风险。"
        )
    if n == "07_一层开关插座布置图.png":
        return (
            "一层开关插座布置反映照明分区、吧台设备与顾客可及插座的基本逻辑，应与平面功能名称一一对照。"
        )
    if n == "08_二层开关插座布置图.png":
        return (
            "二层插座布置需兼顾散座阅读、围炉区安全距离与造景区维护，阅读时关注高度与家具遮挡关系。"
        )
    if n == "09_一层吊顶布置图.png":
        return (
            "吊顶布置整合灯具、风口与检修口位置，是立面节奏与室内净高控制的直接依据；应与剖面标高互证。"
        )
    if n == "10_二层吊顶布置图.png":
        return (
            "二层吊顶常承担氛围照明与设备隐蔽双重任务，读图需核对围炉区上方是否有足够净高与检修可达性。"
        )

    if n == "设计理念图.png":
        return (
            "理念图以“传承—创新—融合”为骨架展开思维导图：物质层面强调传统建筑语汇与器物的当代表达，非物质层面强调茶事礼仪与场景再现；"
            "创新枝指向饮品文化、空间风格、顾客体验与光影装置；融合枝强调中西并置、业态叠合与五感体验。"
            "圆形实景拼贴呈现木作茶台、竹影墙、现代沙发与器物细节，可视为后续彩平与效果图的“母题拼盘”。"
            "图示标题中“水街”语境与彭州龙兴寺片区开题叙事一致，可作为全案概念对齐的图解依据。"
        )

    # —— 平面与彩平 ——
    if n == "一层平面图.jpg":
        return (
            "一层平面在约14.2m见方柱网内组织入口—LOGO墙—服务前台的主轴序列，左侧布置品茶展示与冲泡体验并结合室内绿化，"
            "中部以圆形铺装与半围合界面形成“第二展示区”，右侧以半圆卡座形成雅座区，卫生间与楼梯等功能集中于边角。"
            "曲面隔墙与家具界定分区而非实墙满隔，使视线可穿透，符合茶空间“通透而有层次”的体验目标。"
        )
    if n == "一层彩平图.jpg":
        return (
            "彩平在一层平面基础上以材质色域区分地面铺装与家具组团，读图重点在于主通道亮度、展示区焦点与座位区的围合感。"
            "色彩梯度应与立面与灯光策略一致，避免出现“平面很满、立面很空”的断裂。"
        )
    if n == "二层平面图.jpg":
        return (
            "二层平面延续开题所述“散座+造景+围炉+雅间”组合逻辑，楼梯位置决定竖向联系与后场补给路径。"
            "阅读时应核对二层是否形成相对安静的内区，以及围炉与植物造景是否占用疏散宽度。"
        )
    if n == "二层彩平图.jpg":
        return (
            "二层彩平强调体验区与社交区的氛围差异，可通过冷暖对比与地毯边界提示停留时长不同的空间类型。"
        )

    # —— 流线与功能分区 ——
    if n == "一层交通流线图.jpg":
        return (
            "一层交通流线图以颜色区分品茶展示、点单、散座、茶咖展示、雅座、冲泡体验、前台与卫生间等功能，"
            "红色粗虚线表达主入口—前台—散座核—点单区的主要动线，蓝色细虚线表达次级参观与体验回路。"
            "读图可见设计试图建立“先看见—再体验—再消费”的行为节奏，冲泡体验区相对独立又可通过次级流线被自然发现。"
        )
    if n == "一层功能分区图.jpg":
        return (
            "一层功能分区图以面域涂色表达各业态板块边界，可与交通流线图对照检查是否存在动线穿越静区或后勤穿越顾客核心区的问题。"
        )
    if n == "二层交通流线图.jpg":
        return (
            "二层交通流线强调上楼后的第一眼组织：散座与造景是否形成缓冲，围炉区是否成为视觉与行为焦点，雅间是否获得相对独立走廊。"
        )
    if n == "二层功能分区图.jpg":
        return (
            "二层功能分区图用于校核体验型业态（围炉、冲泡、仓储）与卫生、楼梯之间的关系，避免功能拼贴导致运营交叉。"
        )

    # —— 水电 ——
    if n == "一楼水电分析图.png":
        return (
            "一层水电分析标示下水点位与可布置洁具、清洗区的潜力，是卫生间与吧台选址的硬条件。"
            "读图应结合平面核对排水坡度与噪音区是否远离安静雅座。"
        )
    if n == "二楼水电分析图.png":
        return (
            "二层下水点位数量少于一层，阅读重点在于是否满足冲泡区清洗与卫生间布置，必要时需采用提升或横管组织方案（以施工图为准）。"
        )

    # —— 立面 ——
    if n == "正立面图.jpg":
        return (
            "正立面呈现深色大挑檐与竖向木百叶的节奏变化，首层入口玻璃门与人物剪影建立尺度参照。"
            "材料以暖木、浅色实墙与深色屋面形成横向分层，体现“屋檐语汇+当代体块”的并置，为室内材料向外延伸提供线索。"
        )
    if n == "后立面图.jpg":
        return (
            "后立面阅读重点在于服务与后勤开口是否隐蔽、设备百叶是否破坏整体节奏，以及与相邻建筑的界面关系。"
        )
    if n == "侧立面图.jpg":
        return (
            "侧立面揭示建筑进深与屋面坡度在侧面观看时的比例，可辅助判断二层采光与临街展示面的可达性。"
        )

    # —— 剖面 ——
    if n == "剖面图.png":
        return (
            "剖面给出主要标高与屋面坡度约25°的关系，显示二层净空大于一层，楼梯为折返形式。"
            "该图是吊顶分层、灯具安装高度与空调风口布置的基准，也可用于核对开题所述层高数据与空间尺度叙述。"
        )
    if n.startswith("剖面图") and n.endswith(".jpg"):
        return (
            "补充剖面从不同切割位置揭示局部净高与楼梯平台关系，阅读时应关注是否出现碰头线、梁下净高不足或栏杆可视性问题。"
        )

    # —— 效果图（逐张简析，强调空间氛围与使用情境）——
    if parent == "07_效果图":
        scenes = {
            "休息区效果图.jpg": "休息区强调围合与织物触感，灯光偏暖，适合作为社交强度较低、停留时间中等的节点。",
            "冲泡区效果图.jpg": "冲泡区强调操作台可视性与器物陈列，是“茶文化可读性”的关键场景，应注意与顾客动线的观看距离。",
            "制茶体验区效果图.jpg": "制茶体验区突出参与感与展示面，家具与照明应服务于操作演示而非遮挡视线。",
            "前台效果图.jpg": "前台作为第一印象界面，材料与LOGO墙形成品牌识别，需与入口地面铺装及排队空间协同。",
            "卡座效果图.jpg": "卡座区强调半私密与尺度亲切，隔断高度与声环境密切相关。",
            "围炉煮茶区效果图.jpg": "围炉区具仪式感与热源集中特征，读图需关注座椅围合、排烟与安全的视觉提示。",
            "大厅效果图.jpg": "大厅夜景效果以大面积玻璃借景街道，室内木格栅天花与室内造景树形成“内外互换”的视觉层次，体现静谧商业氛围。",
            "散座效果图.jpg": "散座强调灵活与通透，适合作为二层“呼吸带”的主体。",
            "散座效果图1.jpg": "散座变体角度补充表达家具组合与灯具关系，可验证平面模块在三维中的一致性。",
            "楼梯间效果图.jpg": "楼梯间作为竖向交通形象面，材料与照明影响连续体验，应避免仅作“剩余空间”。",
        }
        return scenes.get(n, f"《{fig.stem}》从氛围、材料与家具尺度三方面体现该功能节点的设计意图，可与同区平面图点位互证。")

    # —— 轴测与景观 ——
    if n == "一楼轴测图.jpg":
        return (
            "一层轴测帮助建立家具群组与曲面隔断的三维关系，可检验平面上的“开放”是否在高度方向出现压抑或杂乱。"
        )
    if n == "二楼轴测图.jpg":
        return (
            "二层轴测突出楼梯、散座与造景之间的咬合，阅读重点在于视线是否通透、围炉区是否成为空间锚点。"
        )
    if n == "二楼景观.jpg":
        return (
            "二层景观视角强调外部界面与室内看景关系，可支持论文关于“借景—框景”的论述，并提示玻璃反光与隐私控制。"
        )

    return (
        f"《{fig.stem}》归类于「{parent}」。结合开题对彭州龙兴寺水街茶咖的定位，"
        f"该图纸用于支撑相应章节的空间论证；阅读时应将图例、标注与平面索引对照，避免脱离尺度与功能谈形式。"
    )


def figures_for_chapter(heading_line: str) -> list[Path]:
    """按章节返回精选随文插图：前期分析、方案功能、技术深化；不含效果图。"""
    h = heading_line.strip()
    if "第3章" in h:
        return figures_in_order(
            "01_前期与综合分析",
            [
                "区位分析图.png",
                "场地交通流线分析.png",
                "场地功能分区分析.png",
                "建筑结构分析图.png",
                "冬夏两季日照分析.png",
                "通风分析图.png",
            ],
        )
    if "第5章" in h:
        return figures_from_tuples(
            [
                ("02_平面与彩平图", "一层平面图.jpg"),
                ("02_平面与彩平图", "二层平面图.jpg"),
                ("02_平面与彩平图", "一层彩平图.jpg"),
                ("03_流线与功能分区图", "一层交通流线图.jpg"),
                ("03_流线与功能分区图", "一层功能分区图.jpg"),
                ("03_流线与功能分区图", "二层交通流线图.jpg"),
                ("03_流线与功能分区图", "二层功能分区图.jpg"),
                ("04_水电分析图", "一楼水电分析图.png"),
                ("04_水电分析图", "二楼水电分析图.png"),
                ("05_立面图", "正立面图.jpg"),
                ("06_剖面图", "剖面图.png"),
            ]
        )
    if "第6章" in h:
        return figures_in_order(
            "09_PDF与技术图纸导出",
            [
                "03_一层建筑墙体改造图.png",
                "04_二层建筑墙体改造图.png",
                "05_一层地面铺装图.png",
                "06_二层地面铺装图.png",
                "07_一层开关插座布置图.png",
                "08_二层开关插座布置图.png",
                "09_一层吊顶布置图.png",
                "10_二层吊顶布置图.png",
            ],
        )
    return []


def strip_leading_book_title(md: str) -> str:
    """去掉文件首行 `# 题名`，保留从第一个 `## ` 起的内容用于按章拆分。"""
    md = md.strip()
    ls = md.splitlines()
    i = 0
    if i < len(ls) and ls[i].startswith("# ") and not ls[i].startswith("##"):
        i += 1
    while i < len(ls) and ls[i].strip() == "":
        i += 1
    return "\n".join(ls[i:]).strip()


def split_chapters(md: str) -> list[tuple[str, str]]:
    rest = strip_leading_book_title(md)
    parts = re.split(r"(?m)^(?=## )", rest)
    chapters: list[tuple[str, str]] = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        ls = part.splitlines()
        head = ls[0].strip()
        body = "\n".join(ls[1:]).strip()
        chapters.append((head, body))
    return chapters


def split_body_to_blocks(body: str) -> list[tuple[str, str]]:
    """`##`→一级标题，`###`→二级标题，段落按空行合并。"""
    blocks: list[tuple[str, str]] = []
    buf: list[str] = []

    def flush_para():
        nonlocal buf
        if buf:
            blocks.append(("p", "\n".join(buf).strip()))
            buf = []

    for line in body.splitlines():
        if line.startswith("### "):
            flush_para()
            blocks.append(("h2", line[4:].strip()))
        elif line.startswith("## "):
            flush_para()
            blocks.append(("h1", line[3:].strip()))
        elif line.strip() == "":
            flush_para()
        else:
            buf.append(line.rstrip())
    flush_para()
    return blocks


def add_body_blocks(doc: Document, blocks: list[tuple[str, str]], *, references_mode: bool = False):
    """正文段落用 Normal（宋体小四 + 1.5 倍行距 + 首行缩进）；参考文献条目用五号。"""
    body_pt = 10.5 if references_mode else 12
    for kind, text in blocks:
        text = text.strip()
        if not text:
            continue
        if kind == "h1":
            p = doc.add_heading(text, level=1)
            for r in p.runs:
                set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size_pt=15, bold=True)
        elif kind == "h2":
            p = doc.add_heading(text, level=2)
            for r in p.runs:
                set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size_pt=14, bold=True)
        else:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            pf.first_line_indent = Cm(0.74)
            paragraph_add_body_runs(p, text, body_pt=body_pt)


def add_figure_block_docx(doc: Document, fig_no: int, fig: Path, analysis: str):
    """
    规范：插图随文；图题置于图下方，图号与图名之间空一格，黑体五号字。
    顺序：读图分析（正文）→ 插图（居中）→ 图题（居中）。
    段落间距与分页属性收紧，减少页底大块留白。
    """
    pa = doc.add_paragraph()
    pa.style = doc.styles["Normal"]
    pa.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pa.paragraph_format.line_spacing = 1.5
    pa.paragraph_format.first_line_indent = Cm(0.74)
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after = Pt(3)
    pa.paragraph_format.keep_together = False
    pa.paragraph_format.keep_with_next = False
    pa.paragraph_format.widow_control = False
    paragraph_add_body_runs(pa, analysis, body_pt=12)

    picp = doc.add_paragraph()
    picp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picp.paragraph_format.space_before = Pt(2)
    picp.paragraph_format.space_after = Pt(2)
    picp.paragraph_format.first_line_indent = Pt(0)
    picp.paragraph_format.keep_together = False
    picp.paragraph_format.keep_with_next = False
    picp.paragraph_format.widow_control = False
    try:
        picp.add_run().add_picture(str(fig), width=Inches(5.45))
    except Exception:
        err = picp.add_run(f"（插图载入失败：{fig.name}）")
        set_run_font(err, east_asia="宋体", size_pt=12)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_together = False
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.widow_control = False
    cr = cap.add_run(f"图{fig_no}\u3000{fig.stem}")
    set_run_font(cr, east_asia="黑体", ascii_font="Times New Roman", size_pt=10.5, bold=True)


def build_markdown(body_full: str) -> str:
    book_title = ""
    ls0 = body_full.strip().splitlines()
    if ls0 and ls0[0].startswith("# ") and not ls0[0].startswith("##"):
        book_title = ls0[0][2:].strip()

    parts: list[str] = []
    if book_title:
        parts.append(f"# {book_title}\n")

    chapters = split_chapters(body_full)
    fig_no = 1
    for head, body in chapters:
        parts.append(f"\n{head}\n\n{body}\n")
        flist = figures_for_chapter(head)
        if not flist:
            continue
        parts.append(f"\n#### 附图\n\n")
        for fig in flist:
            ana = narrative_for_figure(fig)
            parts.append(f"#### 读图分析（图{fig_no}）\n\n{ana}\n\n")
            parts.append(f"![图{fig_no} {fig.stem}]({rel_md_path(fig)})\n\n")
            parts.append(f"**图{fig_no}　{fig.stem}**\n\n")
            fig_no += 1

    return "".join(parts).strip() + "\n"


def thesis_body_plain() -> str:
    p = OUT_DIR / "thesis_body_content.md"
    if not p.is_file():
        raise FileNotFoundError(f"missing {p}")
    return p.read_text(encoding="utf-8").strip()


def try_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return pdf_path.is_file() and pdf_path.stat().st_size > 1000
    except Exception as e:
        print("docx2pdf failed:", e, file=sys.stderr)
        return False


def main() -> None:
    if not DRAWINGS.is_dir():
        sys.exit("missing drawings folder")

    body = thesis_body_plain()
    md_text = build_markdown(body)
    (OUT_DIR / MD_NAME).write_text(md_text, encoding="utf-8")
    print("wrote", OUT_DIR / MD_NAME)

    doc = Document()
    configure_document_styles(doc)

    sect = doc.sections[0]
    sect.top_margin = Cm(2.5)
    sect.bottom_margin = Cm(2)
    sect.left_margin = Cm(2.5)
    sect.right_margin = Cm(2)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.first_line_indent = Pt(0)
    tp.paragraph_format.space_after = Pt(12)
    tr = tp.add_run("彭州市水街茶咖啡厅室内设计")
    set_run_font(tr, east_asia="黑体", ascii_font="Times New Roman", size_pt=18, bold=True)

    sub = doc.add_paragraph("本科毕业论文（设计）·环境设计专业\n学生：李青桐　学号：202201997　学院：风景园林学院　导师：张益昇（讲师）")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.first_line_indent = Pt(0)
    sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    sub.paragraph_format.line_spacing = 1.5
    for r in sub.runs:
        set_run_font(r, east_asia="宋体", ascii_font="Times New Roman", size_pt=12)

    chapters = split_chapters(body)
    fig_no = 1
    for head, body in chapters:
        hh = doc.add_heading(head.replace("## ", "").strip(), level=1)
        for r in hh.runs:
            set_run_font(r, east_asia="黑体", ascii_font="Times New Roman", size_pt=15, bold=True)
        hh.paragraph_format.page_break_before = False
        ref_mode = "参考文献" in head
        add_body_blocks(doc, split_body_to_blocks(body), references_mode=ref_mode)
        flist = figures_for_chapter(head)
        for fig in flist:
            add_figure_block_docx(doc, fig_no, fig, narrative_for_figure(fig))
            fig_no += 1

    docx_path = OUT_DIR / DOCX_NAME
    doc.save(str(docx_path))
    print("wrote", docx_path)

    pdf_path = OUT_DIR / PDF_NAME
    if try_docx_to_pdf(docx_path, pdf_path):
        print("wrote", pdf_path)
    else:
        print("PDF export failed", file=sys.stderr)


if __name__ == "__main__":
    main()
